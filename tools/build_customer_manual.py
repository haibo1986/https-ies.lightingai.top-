from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

OUT = Path(__file__).resolve().parents[1] / "IES光度文件快速换算工具_客户使用说明书.docx"
doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.27), Inches(11.69)
sec.top_margin, sec.bottom_margin = Inches(.72), Inches(.68)
sec.left_margin, sec.right_margin = Inches(.78), Inches(.78)

GREEN = "176B55"; DARK = "20322B"; LIME = "B8D33A"; PALE = "EAF3EE"; GOLD = "B87915"; RED = "A43A2A"; GRAY="5E6A64"

styles = doc.styles
normal = styles['Normal']; normal.font.name='Microsoft YaHei'; normal.font.size=Pt(10)
normal._element.rPr.rFonts.set(qn('w:eastAsia'),'Microsoft YaHei')
normal.paragraph_format.space_after=Pt(5); normal.paragraph_format.line_spacing=1.25
for name,size,color,before,after in [('Title',28,DARK,0,10),('Heading 1',18,GREEN,16,7),('Heading 2',13,DARK,11,5),('Heading 3',11,GREEN,8,3)]:
    st=styles[name]; st.font.name='Microsoft YaHei'; st._element.rPr.rFonts.set(qn('w:eastAsia'),'Microsoft YaHei'); st.font.size=Pt(size); st.font.bold=True; st.font.color.rgb=RGBColor.from_string(color); st.paragraph_format.space_before=Pt(before); st.paragraph_format.space_after=Pt(after); st.paragraph_format.keep_with_next=True
for name in ['List Bullet','List Number']:
    st=styles[name]; st.font.name='Microsoft YaHei'; st._element.rPr.rFonts.set(qn('w:eastAsia'),'Microsoft YaHei'); st.font.size=Pt(10); st.paragraph_format.space_after=Pt(3)

def shade(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),fill); tcPr.append(shd)
def margins(cell, top=100, start=130, bottom=100, end=130):
    tc=cell._tc.get_or_add_tcPr(); mar=tc.first_child_found_in('w:tcMar')
    if mar is None: mar=OxmlElement('w:tcMar'); tc.append(mar)
    for side,val in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        el=OxmlElement('w:'+side); el.set(qn('w:w'),str(val)); el.set(qn('w:type'),'dxa'); mar.append(el)
def set_cell_text(cell,text,bold=False,color=None,size=9):
    cell.text=''; p=cell.paragraphs[0]; p.paragraph_format.space_after=Pt(0); r=p.add_run(str(text)); r.bold=bold; r.font.name='Microsoft YaHei'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'Microsoft YaHei'); r.font.size=Pt(size)
    if color:r.font.color.rgb=RGBColor.from_string(color)
    cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; margins(cell)
def table(headers, rows, widths=None):
    t=doc.add_table(rows=1, cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    for i,h in enumerate(headers): set_cell_text(t.rows[0].cells[i],h,True,'FFFFFF',9); shade(t.rows[0].cells[i],GREEN)
    for ri,row in enumerate(rows):
        cells=t.add_row().cells
        for i,v in enumerate(row): set_cell_text(cells[i],v,False,DARK,8.7); shade(cells[i], 'F5F8F6' if ri%2 else 'FFFFFF')
    if widths:
        for row in t.rows:
            for c,w in zip(row.cells,widths): c.width=Inches(w)
    doc.add_paragraph().paragraph_format.space_after=Pt(1)
    return t
def callout(label,text,kind='info'):
    colors={'info':(PALE,GREEN),'warn':('FFF3D8',GOLD),'danger':('FBE7E3',RED)}; fill,color=colors[kind]
    t=doc.add_table(rows=1,cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER; c=t.cell(0,0); shade(c,fill); margins(c,150,170,150,170)
    p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(0); r=p.add_run(label+'  '); r.bold=True; r.font.color.rgb=RGBColor.from_string(color); r2=p.add_run(text); r2.font.color.rgb=RGBColor.from_string(DARK)
    doc.add_paragraph().paragraph_format.space_after=Pt(0)
def bullets(items, numbered=False):
    for x in items: doc.add_paragraph(x, style='List Number' if numbered else 'List Bullet')
def step(title, body):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(7); p.paragraph_format.space_after=Pt(3); r=p.add_run(title); r.bold=True; r.font.size=Pt(11); r.font.color.rgb=RGBColor.from_string(GREEN)
    p=doc.add_paragraph(body); p.paragraph_format.left_indent=Inches(.12)
def page(): doc.add_page_break()

# Cover
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(72)
r=p.add_run('IES'); r.bold=True; r.font.size=Pt(42); r.font.color.rgb=RGBColor.from_string(GREEN)
p=doc.add_paragraph('PHOTOMETRIC TOOLS'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.runs[0].font.size=Pt(10); p.runs[0].font.color.rgb=RGBColor.from_string(GRAY)
p=doc.add_paragraph('IES 光度文件快速换算工具',style='Title'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(48)
p=doc.add_paragraph('客户使用说明书'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.runs[0].bold=True; p.runs[0].font.size=Pt(18); p.runs[0].font.color.rgb=RGBColor.from_string(GREEN)
p=doc.add_paragraph('适用于当前软件版本 · 2026 年 8 月'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(15); p.runs[0].font.color.rgb=RGBColor.from_string(GRAY)
doc.add_paragraph('\n')
callout('请先阅读', '本软件生成的是估算版 IES 与专业光度报告，仅用于方案模拟、内部评估和客户初步沟通，不能替代实验室实测、认证、验收或法律依据。','warn')
p=doc.add_paragraph('建议阅读顺序：快速上手 → 参数填写 → 结果下载 → 常见问题'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(30); p.runs[0].font.color.rgb=RGBColor.from_string(GRAY)

page(); doc.add_heading('1. 软件能做什么',0)
doc.add_paragraph('本软件以一份已实测的 IES 光度文件为基础，在配光形状保持不变的前提下，按“目标光通量 ÷ 原始实测总光通量”的比例缩放光强矩阵，并生成目标型号的估算交付文件。')
doc.add_heading('主要功能',1)
table(['功能','客户可获得的结果'],[
('IES 文件解析','识别标准版本、TILT、功率、光通量、角度数量、最大光强、灯具尺寸及文件标签。'),
('光度可视化','显示光强分布曲线、主要平面、50% 光束角和 10% 光场角等摘要。'),
('目标参数换算','按目标光通量比例缩放 candela 光强数据，并写入目标型号、功率和发光尺寸。'),
('目标光通量估算','可按 LED 电流曲线、LED 颗数、模组数量、PWM 占空比或功率×光效辅助估算。'),
('风险判断','低/中风险可生成；光束角、透镜或光学结构变化会被阻止并要求重新实测。'),
('专业报告','生成统一版式 PDF，并可预览；另提供 HTML 与 Markdown 辅助文件。'),
('结果校验','检查生成 IES、光强矩阵、功率、光通量、发光面尺寸、PDF 内容等关键项目。'),
], [1.45,5.05])
doc.add_heading('适用场景',1)
bullets(['同系列灯具，仅功率或驱动电流变化。','LED 数量或密度变化，但透镜、角度和光学结构不变。','灯具长度或模组数量变化，用于前期方案评估。','PWM 调光或缺少目标光通量时的辅助估算。'])
doc.add_heading('不适用场景',1)
bullets(['光束角变化、透镜更换、反光杯或其他光学结构变化。','需要认证、招投标定版、工程验收、第三方检测或法律用途。','需要 DIALux、Relux、AGi32 场景模拟，或需要批量数据库管理。'])

page(); doc.add_heading('2. 使用前准备',0)
doc.add_heading('您需要准备的资料',1)
table(['资料','要求','是否必需'],[
('原始 IES 文件','扩展名 .ies；不超过 10 MB；LM-63-1995 或 LM-63-2002；TILT=NONE。','必需'),
('原始实测总光通量','单位 lm。绝对光度文件必须人工提供。','必需'),
('目标型号、功率、光通量','功率单位 W，光通量单位 lm，数值均须大于 0。','必需'),
('目标发光面尺寸','长度×宽度，单位 mm，用于亮度计算并写入 IES。','必需'),
('原始测试 PDF','仅作溯源附件和信息核对，不作为新报告模板。','可选'),
('企业与产品资料','Logo、企业名称、制造商、电气/颜色/外形参数、备注。','可选'),
], [1.3,4.1,1.1])
callout('文件兼容性', '软件目前只接受 LM-63-1995、LM-63-2002 且 TILT=NONE 的 IES。TILT=INCLUDE、TILT=FILE 或损坏/非标准文件不会被自动修复。','info')
doc.add_heading('进入软件',1)
doc.add_paragraph('请使用软件管理员提供的网址，在 Chrome 或 Edge 浏览器中打开。若页面能打开但无法上传，请联系管理员检查后端服务。普通客户不需要安装 Python 或 Node.js。')
doc.add_heading('界面流程',1)
bullets(['步骤 1：上传原始 IES。','步骤 2：检查解析结果并填写目标参数。','步骤 3：生成估算 IES 与专业光度报告。','结果中心：预览、校验并下载文件。'], numbered=True)

page(); doc.add_heading('3. 快速上手：完成一次换算',0)
step('步骤 1｜上传并解析 IES','点击“选择 .ies 文件”，选中原始实测文件，再点击“上传并解析”。等待页面显示 IES 光度摘要。')
step('步骤 2｜核对解析结果','重点检查原始文件名、标准版本、输入功率、建议光通量、最大光强、垂直/水平角度数量和配光曲线。如绝对光度文件提示“需手动填写光通量”，请以实测报告为准填写。')
step('步骤 3｜填写必填参数','填写原始实测总光通量、目标型号、目标功率、目标光通量、发光面长度和宽度。所有数字必须大于 0。')
step('步骤 4｜选择变更类型','按实际情况选择“仅功率或电流变化”“LED 数量或密度变化”或“灯具长度或模组变化”。看到黄色中风险提示时，可继续用于初步模拟，但需关注均匀性。')
step('步骤 5｜补充报告信息（可选）','可上传原始光度测试 PDF，并填写企业名称、Logo、制造商、报告编号、日期、电气参数、色温、显色指数、灯具外形、计算条件和备注。')
step('步骤 6｜生成','点击“生成估算 IES + 专业光度报告”。处理中请不要刷新页面或关闭浏览器。')
step('步骤 7｜检查与下载','生成成功后先确认“校验：全部通过”，再到“文件中心”下载 IES 和 PDF；需要网页或内部留档时再下载 HTML、MD。')
callout('关键原则', '换算比例使用光通量比，不使用功率比。LED 光效、驱动效率和温升可能随功率变化，因此功率不能可靠代表光输出。','warn')

page(); doc.add_heading('4. 参数填写说明',0)
table(['字段','填写方法','常见错误'],[
('原始实测总光通量','填写原灯具实际测得的总光通量（lm）。','误填额定值、单颗 LED 光通量或功率。'),
('目标型号','填写最终文件希望使用的产品型号。','包含过多特殊字符；型号与客户订单不一致。'),
('目标功率','目标灯具输入功率（W），必须大于 0。','误填驱动输出功率或电流。'),
('目标光通量','希望换算到的总光通量（lm）。','按功率比例直接猜测，未考虑光效与温升。'),
('发光面长/宽','实际发光口尺寸（mm），用于亮度曲线。','误填整灯外形尺寸；长宽单位写成 m。'),
('变更类型','必须与实际产品变化一致。','为了生成而选择低风险类别。'),
], [1.25,3.25,2.0])
doc.add_heading('风险等级如何理解',1)
table(['等级','场景','软件行为'],[
('低风险','仅功率或电流变化，且同系列、同角度、同光学结构。','允许生成。'),
('中风险','LED 数量/密度或灯具长度/模组变化。','允许生成，并提示验证近场与墙面均匀性。'),
('高风险','光束角、透镜或光学结构变化。','禁止生成，必须重新光度实测。'),
], [1.0,3.7,1.8])
doc.add_heading('报告补充信息',1)
doc.add_paragraph('补充字段属于用户提供数据，会在报告中与 IES 计算数据分开标识。Logo 仅支持 PNG/JPG，最大 2 MB；备注最多 500 字。原始 PDF 的识别结果只用于核对，不会覆盖目标 IES 的计算结果。')

page(); doc.add_heading('5. 不知道目标光通量时',0)
doc.add_paragraph('展开“不知道目标光通量？”估算器，选择与目标灯具最接近的变化方式。估算结果会显示建议范围和置信度；点击“使用这个估算值”后，软件会自动回填目标光通量及相应变更类型。')
table(['估算方式','需要填写','注意事项'],[
('调整 LED 电流','原/目标单颗电流；至少 2 个电流—相对光通量数据点。','仅在数据曲线范围内插值，不会外推。'),
('调整 LED 颗数','原/目标颗数；原/目标发光长度；目标宽度。','软件检查 LED 密度；密度明显下降可能阻止采用。'),
('颗数和电流都变','上述电流曲线及原/目标 LED 颗数。','散热未知时置信度降低。'),
('灯具长度/模组变化','原/目标模组数；目标发光口长宽。','同步更新 IES 尺寸，近场均匀性仍需验证。'),
('PWM 调光','原/目标占空比（0–100%）。','仅适用于峰值电流不变的 PWM 调光。'),
('功率×光效','目标功率及预计目标光效（lm/W）。','低置信度备用方法。'),
], [1.35,3.15,2.0])
doc.add_heading('导入 LED 曲线',1)
bullets(['可直接粘贴数据：每行“电流, 相对光通量”。','支持 Excel（.xlsx/.xls）、CSV 和 TXT；默认读取第一个工作表前两列。','也可上传 PNG/JPG/WebP 曲线截图，标定坐标轴后半自动取点；至少确认 2 个数据点。'])
callout('禁止简单换算', '若透镜或光学结构发生变化，估算器会提示重新实测。仅材料透过率变化时，可输入原/目标透过率作为修正因子。','danger')

page(); doc.add_heading('6. 查看结果与下载文件',0)
doc.add_heading('IES 预览',1)
bullets(['确认目标型号、目标功率、目标光通量和最大光强。','查看“文件检查”，优先确保每项均为 ✓。','核对生成后的配光曲线是否与原文件形状一致，仅幅值按比例变化。','必要时展开 IES 原始文本（页面显示前 120 行）进行技术核对。'])
doc.add_heading('报告预览',1)
doc.add_paragraph('“报告预览”标签中可直接查看 PDF，也可点击“单独打开 PDF”。当前统一专业报告为 13 页，内容包括经典配光页面、亮度限制、等照度、区域光通量及完整光强矩阵等。')
doc.add_heading('文件中心',1)
table(['文件','用途','建议'],[
('目标型号.ies','照明设计软件使用的估算配光文件。','核心交付；文件内带 ESTIMATED 标识。'),
('目标型号-report.pdf','客户阅读的专业估算光度报告。','核心交付；保留估算声明。'),
('HTML 报告','浏览器网页预览。','辅助查看或内部分享。'),
('Markdown 报告','纯文本结构的内部说明。','用于内部留档或二次编辑。'),
('原始测试 PDF','用户上传的来源文件。','只读保留，不被修改。'),
], [1.35,3.05,2.1])
callout('保存文件', '生成文件属于临时运行文件。建议生成后立即下载并按贵司项目文件规则归档；服务重启后上传记录会失效，超过 24 小时的运行文件也可能被清理。','warn')

page(); doc.add_heading('7. 常见问题与处理',0)
table(['现象/提示','原因与处理方法'],[
('“文件格式必须是 .ies”','请选择扩展名为 .ies 的文件，不要直接修改其他文件的后缀。'),
('“仅支持 1995 和 2002”','文件不是 LM-63-1995/2002，请向测试机构索取兼容格式。'),
('“仅支持 TILT=NONE”','当前不支持 INCLUDE/FILE；请让测试机构重新导出。'),
('“需手动填写光通量”','这是绝对光度文件（lumens_per_lamp=-1），请从实测报告填写总光通量。'),
('生成按钮不可用','当前选了高风险变更，或仍有必填项未正确填写。'),
('提示必须重新实测','光束角、透镜或光学结构改变，无法通过统一倍率可靠获得配光。'),
('上传记录不存在/服务已重启','重新上传 IES；如上传过来源 PDF，也需重新上传。'),
('页面可打开但无法上传','后端服务可能未运行，请联系管理员。'),
('报告校验存在异常','不要交付；保存提示内容并联系软件管理员排查。'),
('下载后找不到文件','检查浏览器下载列表或“下载”文件夹，并重新命名归档。'),
], [2.05,4.45])
doc.add_heading('客户提交问题时请提供',1)
bullets(['原始 IES 文件（如可提供）。','操作时填写的目标参数及变更类型。','完整错误提示截图。','浏览器类型、出现时间，以及是否发生过服务重启。'])

page(); doc.add_heading('8. 安全、准确性与交付声明',0)
doc.add_heading('准确性边界',1)
doc.add_paragraph('本软件的核心假设是：目标灯具与原灯具的配光形状保持不变，只对所有光强数据按同一光通量比例缩放。该方法不能预测光学结构变化、热效应、驱动效率变化、灯具长度变化造成的空间均匀性差异，也不能替代真实光度测量。')
doc.add_heading('交付前检查清单',1)
bullets(['原始 IES 来自可信的光度实测。','原始实测总光通量填写正确。','目标型号、功率、光通量和发光面尺寸与产品资料一致。','变更类型与真实产品变化一致，没有规避高风险拦截。','结果页显示校验全部通过。','IES 和 PDF 均已下载、打开并检查。','客户已知文件为 ESTIMATED 估算结果。','正式认证、验收或定版前已安排重新实测。'])
callout('正式声明', '输出仅适用于前期方案模拟、内部评估和客户初步沟通，不是认证级光度测试结果，不得作为正式认证、验收、第三方检测报告或法律争议依据。','danger')
doc.add_heading('管理员附录：本地启动（仅技术人员）',1)
doc.add_paragraph('后端要求 Python 3.10+：进入 backend，创建虚拟环境、安装 requirements.txt，然后运行 uvicorn app.main:app --reload。前端要求 Node.js 20+：进入 frontend，执行 npm install 和 npm run dev。默认访问 http://127.0.0.1:5173；后端健康检查为 http://127.0.0.1:8000/api/health。')
doc.add_paragraph('当前版本验证状态：前端生产构建成功；后端自动化测试 44 项全部通过（验证日期：2026-08-13）。')

# footer + page numbers
for section in doc.sections:
    footer=section.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=footer.add_run('IES 光度文件快速换算工具 · 客户使用说明书  |  '); r.font.size=Pt(8); r.font.color.rgb=RGBColor.from_string(GRAY)
    fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); footer._p.append(fld)

doc.core_properties.title='IES 光度文件快速换算工具客户使用说明书'
doc.core_properties.subject='客户操作、参数说明、结果下载与注意事项'
doc.core_properties.author='IES Photometric Tools'
doc.save(OUT)
print(OUT)
